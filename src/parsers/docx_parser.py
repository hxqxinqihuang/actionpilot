from __future__ import annotations

from io import BytesIO

from docx import Document

from src.parsers.cleaning import clean_extracted_text
from src.parsers.exceptions import CorruptedFileError
from src.parsers.models import ParseResult


def parse_docx(file_name: str, file_bytes: bytes) -> ParseResult:
    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise CorruptedFileError("DOCX could not be opened. The file may be damaged or not a valid DOCX.") from exc

    parts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            if row_text:
                parts.append(row_text)

    text = clean_extracted_text("\n".join(parts))
    return ParseResult(
        text=text,
        file_name=file_name,
        file_type="docx",
        char_count=len(text),
        page_count=None,
        warnings=[],
    )
